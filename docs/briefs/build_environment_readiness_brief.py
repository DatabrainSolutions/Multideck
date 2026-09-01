from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT = "/Users/andrewphillips/Documents/GitHub/Multideck/docs/briefs/multideck-environment-readiness-brief.docx"
LOGO = "/Users/andrewphillips/Documents/GitHub/Multideck/docs/briefs/.assets/multideck-full-logo.svg.png"
COVER = "/Users/andrewphillips/Documents/GitHub/Multideck/docs/briefs/.assets/environment-readiness-cover.png"

INK = "0B1413"
MUTED = "4B5A57"
TEAL = "0E7D74"
DEEP_TEAL = "075D57"
MINT = "E7F1EE"
SOFT = "F5F8F7"
HAIRLINE = "D8E3E0"
CREAM = "FFF7E6"
WHITE = "FFFFFF"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=HAIRLINE, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:color"), color)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for index, width in enumerate(widths_dxa):
        grid.gridCol_lst[index].set(qn("w:w"), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])


def set_run(run, size=12, color=MUTED, bold=False, all_caps=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.font.all_caps = all_caps


def set_para(p, before=0, after=10, line=1.35, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def add_text(parent, text, size=12, color=MUTED, bold=False, before=0, after=10, line=1.35, align=None, all_caps=False):
    p = parent.add_paragraph()
    set_para(p, before, after, line, align)
    set_run(p.add_run(text), size, color, bold, all_caps)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 14}
    p = doc.add_paragraph()
    set_para(p, before=18 if level == 1 else 12, after=7, line=1.1)
    set_run(p.add_run(text), sizes[level], INK, False)
    p.paragraph_format.keep_with_next = True
    return p


def add_kicker(doc, text):
    p = doc.add_paragraph()
    set_para(p, before=0, after=6, line=1.0)
    set_run(p.add_run(text), 10, TEAL, True, True)
    p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, before=0, after=4, line=1.25)
    p.paragraph_format.left_indent = Inches(0.24)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    set_run(p.add_run(text), 11.5, MUTED)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_para(p, before=0, after=5, line=1.25)
    p.paragraph_format.left_indent = Inches(0.27)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    set_run(p.add_run(text), 11.5, MUTED)
    return p


def add_rule(p, color=HAIRLINE):
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "10")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_decision_strip(doc, title, text, caution=False):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9820])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CREAM if caution else DEEP_TEAL)
    set_cell_margins(cell, top=150, start=220, bottom=150, end=220)
    color = INK if caution else WHITE
    add_text(cell, title, size=10, color=color, bold=True, before=0, after=3, line=1.0, all_caps=True)
    add_text(cell, text, size=11.5, color=color, before=0, after=0, line=1.25)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, MINT)
        set_cell_margins(cell)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_para(p, after=0, line=1.05)
        set_run(p.add_run(text), 10.5, TEAL, True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_cell_shading(cell, SOFT)
            set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            set_para(p, after=0, line=1.15)
            set_run(p.add_run(text), 10.2, INK if i == 0 else MUTED, i == 0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_footer(section, label):
    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Pt(491))
    set_table_geometry(table, [7350, 2470])
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    p_left = table.cell(0, 0).paragraphs[0]
    set_para(p_left, after=0, line=1)
    set_run(p_left.add_run(label), 8.5, MUTED, False, all_caps=True)
    p_right = table.cell(0, 1).paragraphs[0]
    set_para(p_right, after=0, line=1, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_run(p_right.add_run("PAGE "), 8.5, MUTED, False)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p_right._p.append(fld)


def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    set_para(p, after=0, line=1, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p.add_run().add_picture(LOGO, width=Pt(115))
    rule = header.add_paragraph()
    set_para(rule, after=0, line=1)
    add_rule(rule)


def page_break(doc):
    doc.add_page_break()


def build_cover_image():
    """Build one full-page cover image to preserve the canonical A4 canvas."""
    scale = 2
    width, height = 1191, 1684
    canvas = Image.new("RGB", (width, height), f"#{MINT}")
    draw = ImageDraw.Draw(canvas)
    draw.rectangle((0, 0, width, round(height * 0.12)), fill=f"#{TEAL}")
    regular = "/System/Library/Fonts/Supplemental/Arial.ttf"
    bold = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
    font_kicker = ImageFont.truetype(bold, 20)
    font_title = ImageFont.truetype(regular, 72)
    font_subtitle = ImageFont.truetype(regular, 31)
    font_panel_label = ImageFont.truetype(bold, 20)
    font_panel = ImageFont.truetype(regular, 24)
    font_reference = ImageFont.truetype(regular, 18)

    capsule = (104, 54, 484, 150)
    draw.rounded_rectangle(capsule, radius=16, fill=f"#{WHITE}")
    logo = Image.open(LOGO).convert("RGBA")
    logo.thumbnail((340, 66))
    canvas.alpha_composite(logo, (capsule[0] + (380 - logo.width) // 2, capsule[1] + (96 - logo.height) // 2)) if canvas.mode == "RGBA" else canvas.paste(logo, (capsule[0] + (380 - logo.width) // 2, capsule[1] + (96 - logo.height) // 2), logo)

    x = 104
    draw.text((x, 430), "ENVIRONMENT READINESS", font=font_kicker, fill=f"#{TEAL}")
    draw.multiline_text((x, 476), "Make Dev a safe\ntenant-equivalent platform", font=font_title, fill=f"#{INK}", spacing=4)
    draw.multiline_text((x, 670), "Current-state audit, completed Dev adoption and the route to a\nreliable Dev → Test → Jenkar release path.", font=font_subtitle, fill=f"#{MUTED}", spacing=8)

    panel = (104, 1040, 1087, 1450)
    draw.rounded_rectangle(panel, radius=18, fill=f"#{WHITE}")
    draw.text((panel[0] + 44, panel[1] + 42), "WHAT THIS BRIEF COVERS", font=font_panel_label, fill=f"#{TEAL}")
    items = [
        "The current environment and isolation gaps",
        "The correct Dev, Test and Jenkar system model",
        "The Dev adoption work now completed in Cloud",
        "Release gates that stop untested changes reaching Jenkar",
    ]
    y = panel[1] + 94
    for item in items:
        draw.ellipse((panel[0] + 48, y + 10, panel[0] + 58, y + 20), fill=f"#{TEAL}")
        draw.text((panel[0] + 76, y), item, font=font_panel, fill=f"#{MUTED}")
        y += 56
    draw.text((104, 1584), "MULTIDECK | INTERNAL DELIVERY BRIEF | 26 AUGUST 2026", font=font_reference, fill=f"#{MUTED}")
    canvas.save(COVER)


def build():
    doc = Document()
    first = doc.sections[0]
    first.page_width = Pt(595.3)
    first.page_height = Pt(841.9)
    first.top_margin = Pt(0)
    first.bottom_margin = Pt(0)
    first.left_margin = Pt(0)
    first.right_margin = Pt(0)
    first.header_distance = Pt(0)
    first.footer_distance = Pt(0)
    if doc.paragraphs:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    # Canonical cover: a single A4 mint canvas avoids renderer-specific table splitting.
    build_cover_image()
    cover_paragraph = doc.add_paragraph()
    set_para(cover_paragraph, before=0, after=0, line=1.0)
    cover_paragraph.add_run().add_picture(COVER, width=Pt(595.3), height=Pt(840))

    # Inner pages.
    inner = doc.add_section(WD_SECTION.NEW_PAGE)
    inner.page_width = Pt(595.3)
    inner.page_height = Pt(841.9)
    inner.top_margin = Pt(44)
    inner.bottom_margin = Pt(42)
    inner.left_margin = Pt(52)
    inner.right_margin = Pt(52)
    inner.header_distance = Pt(18)
    inner.footer_distance = Pt(22)
    inner.header.is_linked_to_previous = False
    inner.footer.is_linked_to_previous = False
    add_header(inner)
    add_footer(inner, "MULTIDECK | ENVIRONMENT READINESS BRIEF")

    add_kicker(doc, "DECISION")
    add_text(doc, "Dev is now modelled as a tenant-equivalent platform environment; Test remains the critical missing boundary.", 29, INK, False, 0, 10, 1.08)
    add_text(doc, "The Dev setup record, Cloud API and accounting connection model are now live. The estate still does not provide a safe Dev → Test → customer release path because Test shares Dev data and several intended hostnames are not yet configured.", 12, MUTED, False, 0, 14)
    add_decision_strip(doc, "NON-NEGOTIABLE", "Dev, Test and Jenkar must never share a Supabase project. App and Live may share the same project only when they belong to the same environment or customer tenant.")
    add_heading(doc, "What we have agreed", 1)
    add_bullet(doc, "dev.multideck.app and dev.multideck.live are the continuous-development environment.")
    add_bullet(doc, "test.multideck.app and test.multideck.live are the longer-lived beta and acceptance environment.")
    add_bullet(doc, "jenkar.multideck.app and jenkar.multideck.live are the only customer production tenant surfaces today.")
    add_bullet(doc, "multideck.cloud is a separate control plane. It records safe infrastructure evidence and never stores tenant runtime credentials.")
    add_heading(doc, "Why this matters now", 1)
    add_text(doc, "Accounting settings exposed the underlying issue: Cloud could not reliably describe or verify the environment attached to a workspace. That is a platform readiness problem, not only an accounting problem. The Cloud setup schema and Dev adoption record now remove that immediate blocker.")
    add_decision_strip(doc, "RELEASE RULE", "A feature must work in Dev, then pass a repeatable Test check, before it can be released to Jenkar.", caution=True)

    # Let the page system choose the natural break; forcing a break here can
    # leave an isolated decision strip on the preceding page.
    add_kicker(doc, "CURRENT-STATE AUDIT")
    add_text(doc, "What is live today", 29, INK, False, 0, 10, 1.08)
    add_text(doc, "Read-only checks completed on 26 August 2026. “Reachable” confirms a successful public HTTPS response; it does not confirm that the environment is production-ready.", 12, MUTED, False, 0, 14)
    add_matrix(doc,
        ["Surface", "Observed state", "Readiness"],
        [
            ["dev.multideck.app", "Reachable. Dev branch; HTTP 200 verified.", "Adopted in Cloud as a live platform environment on project aqtwypsuijxlnvtxpuxe."],
            ["test.multideck.app", "Reachable. Deployed from the Test branch.", "Uses the same Supabase project as Dev — not isolated."],
            ["jenkar.multideck.app", "Not publicly resolvable.", "No live Jenkar operator workspace is available."],
            ["dev.multideck.live", "Not publicly resolvable.", "A Dev branch deployment exists, but the custom DNS route is not configured."],
            ["test.multideck.live", "Not publicly resolvable.", "No Test deployment or custom route is configured."],
            ["jenkar.multideck.live", "Reachable.", "Has its own Vercel project and Jenkar Supabase project."],
            ["multideck.cloud", "Reachable. Cloud setup API deployed.", "Dev has a completed adoption run, safe provider records and a live setup checklist."],
        ],
        [2650, 3300, 3870],
    )
    add_heading(doc, "The immediate risk", 1)
    add_text(doc, "Test changes can currently read and write the same data as Dev. This prevents reliable beta validation, makes failures difficult to reproduce, and creates a route for test configuration to affect active development work. It is the highest-priority environment risk.")
    add_decision_strip(doc, "ACCOUNTING IMPACT", "ERPNext and Sage 50 setup must not be tested against a customer tenant until the Test environment has its own project, sandbox provider settings and recorded verification evidence.", caution=True)

    # Continue naturally to preserve the current-state callout with its audit.
    add_kicker(doc, "TARGET ARCHITECTURE")
    add_text(doc, "One isolated project per system", 29, INK, False, 0, 10, 1.08)
    add_text(doc, "Each row is an independent security and delivery boundary. The two product surfaces in a row share that row’s project because they operate on the same environment data.", 12, MUTED, False, 0, 14)
    add_matrix(doc,
        ["System", "Supabase project", "Product surfaces", "Purpose"],
        [
            ["Dev", "Dedicated Development project", "dev.multideck.app\ndev.multideck.live", "Continuous feature development and sandbox integrations."],
            ["Test", "New dedicated Test project", "test.multideck.app\ntest.multideck.live", "Stable beta testing, repeatable checks and release acceptance."],
            ["Jenkar", "Existing dedicated Jenkar project", "jenkar.multideck.app\njenkar.multideck.live", "Customer production operations and customer portal."],
            ["Cloud", "Separate control-plane project", "multideck.cloud", "Safe records, setup evidence, release controls and health checks."],
        ],
        [1450, 2500, 3000, 2870],
    )
    add_heading(doc, "Operating rules", 1)
    add_bullet(doc, "No customer data, Auth users, Storage objects or provider credentials are copied between the four projects.")
    add_bullet(doc, "Each project has its own exact auth origins, redirect allow list, Edge Function secrets and health checks.")
    add_bullet(doc, "Dev and Test use synthetic or approved sandbox data only. Customer provider credentials remain inside the customer’s own project.")
    add_bullet(doc, "Cloud records project references, deployment evidence and safe connection status — never API secrets, database passwords or service-role keys.")

    # Continue naturally to avoid an empty spacer page.
    add_kicker(doc, "GAPS TO CLOSE")
    add_text(doc, "What the platform needs next", 29, INK, False, 0, 10, 1.08)
    add_heading(doc, "1. Establish the missing Test boundary", 1)
    add_number(doc, "Create the dedicated Test Supabase project and apply the approved application baseline, migrations, storage policies and Edge Functions.")
    add_number(doc, "Attach test.multideck.app and test.multideck.live to that project with Test-only public configuration and secrets.")
    add_number(doc, "Seed controlled test data and integration sandboxes. Do not clone Jenkar data or identities.")
    add_heading(doc, "2. Complete the product surfaces", 1)
    add_number(doc, "Configure dev.multideck.live and test.multideck.live DNS and deploy their branch-specific builds.")
    add_number(doc, "Provision jenkar.multideck.app against the existing Jenkar project and verify its exact customer hostname.")
    add_number(doc, "Keep the App and Live build configuration separate per environment, even where they share that environment’s Supabase project.")
    add_heading(doc, "3. Repair Cloud provisioning", 1)
    add_number(doc, "Maintain the deployed guided setup schema, setup steps, provider connection records and safe verification API in the live Cloud project.")
    add_number(doc, "Classify Test as a platform environment when it is created; retain Dev as a platform environment and Jenkar as the customer tenant. Do not bill or provision Dev or Test as customers.")
    add_number(doc, "Use the completed Dev adoption pattern for existing resources; use normal provisioning only for new customer tenants.")
    add_decision_strip(doc, "CONTROL-PLANE RULE", "A green health check alone is not provisioning evidence. Cloud must show the project, deployment, hostname, baseline and verification state for each system.")

    # Continue naturally to preserve the delivery matrix with its gates.
    add_kicker(doc, "DELIVERY PLAN")
    add_text(doc, "A staged route to reliable releases", 29, INK, False, 0, 10, 1.08)
    add_matrix(doc,
        ["Phase", "Outcome", "Accountable role", "Exit check"],
        [
            ["0. Freeze the model", "Document Dev, Test, Jenkar and Cloud as separate system types.", "Product + platform", "No ambiguity over which hostname or project owns a change."],
            ["1. Isolate Test", "Dedicated Test project, App and Live deployments, DNS and sandbox data.", "Platform engineering", "Test can operate without reaching Dev or Jenkar resources."],
            ["2. Complete Jenkar App", "Jenkar operator workspace is deployed on the existing Jenkar project.", "Platform + App", "jenkar.multideck.app signs in and cannot access other projects."],
            ["3. Enable Cloud setup", "Guided setup data model, adoption path and environment/tenant classification are live for Dev.", "Cloud engineering", "Completed: Cloud now shows Dev evidence instead of the generic accounting settings error."],
            ["4. Prove the release gate", "One feature passes Dev then Test before Jenkar release.", "Delivery lead", "Recorded acceptance evidence and clean rollback path."],
        ],
        [1500, 3300, 2200, 2820],
    )
    add_heading(doc, "Release acceptance gates", 1)
    add_bullet(doc, "The same feature build is healthy on both Dev product surfaces.")
    add_bullet(doc, "Test has passed functional, permission, integration, language and right-to-left checks against its own project.")
    add_bullet(doc, "Cloud shows an approved deployment and health evidence for the target Jenkar system.")
    add_bullet(doc, "The Jenkar release uses only Jenkar configuration and has a tested rollback route.")
    add_heading(doc, "First proof case: Accounting", 1)
    add_text(doc, "Use the Accounting settings route as the first end-to-end proof. It now opens against the Dev setup model, records provider connection metadata only, and keeps ERPNext and Sage 50 secrets inside the Dev project. Repeat the same proof in the new Test project before allowing a Jenkar release.")

    # Continue naturally to preserve the proof case with its decision page.
    add_kicker(doc, "NEXT DECISIONS")
    add_text(doc, "What we need to authorise", 29, INK, False, 0, 10, 1.08)
    add_text(doc, "The remaining work changes cloud projects, deployment configuration and DNS. It should be carried out deliberately, in this order, with the existing systems adopted rather than rebuilt.", 12, MUTED, False, 0, 14)
    add_matrix(doc,
        ["Decision", "Recommendation"],
        [
            ["Development project", "Keep the existing Development project as Dev. This is now recorded as a platform environment in Cloud."],
            ["Test project", "Create a new, separate project. Do not reuse Dev or Jenkar."],
            ["Jenkar project", "Keep the existing Jenkar project. Attach both Jenkar App and Jenkar Live to it after verification."],
            ["Cloud data model", "Keep the explicit platform-environment versus customer-tenant classification now deployed for Dev adoption and new provisioning."],
            ["Provider connections", "Use sandbox ERPNext and Sage 50/Hyperext settings in Dev and Test. Store only safe status in Cloud."],
        ],
        [2900, 6920],
    )
    add_decision_strip(doc, "RECOMMENDED NEXT STEP", "Create the isolated Test project and its two hostnames first. Then prove the accounting route in Dev and Test before releasing it to Jenkar.")
    add_heading(doc, "Audit scope and assumptions", 1)
    add_bullet(doc, "This brief is based on the 26 August 2026 audit plus the completed Dev adoption changes to the live Cloud project and Cloud API.")
    add_bullet(doc, "A reachable route is not treated as proof of correct Auth, provider secret, migration or tenant-isolation configuration.")
    add_bullet(doc, "The completed adoption changes recorded safe identifiers, status, checklist and roster evidence only. No provider credentials, tenant data, Auth users, DNS records or deployments were copied or recreated.")

    doc.core_properties.title = "Multideck Environment Readiness Brief"
    doc.core_properties.subject = "Dev, Test and Jenkar environment architecture"
    doc.core_properties.author = "Multideck"
    doc.save(OUT)


if __name__ == "__main__":
    build()

from pathlib import Path
from copy import deepcopy
from shutil import copyfile
from zipfile import ZipFile, ZIP_DEFLATED
from lxml import etree
from docx import Document

REF = Path('/Users/harryphillips/.codex/skills/artifact-template-multideck-calculator-guide/assets/reference.docx')
OUT = Path('/Users/harryphillips/Databrain/GitHub/Multideck/outputs/Multideck_Tenant_Isolation_and_Releases.docx')
WORK = Path('/Users/harryphillips/Databrain/GitHub/Multideck/work/tenant_release_template/working.docx')
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def set_text(paragraph, value):
    if not paragraph.runs:
        paragraph.add_run(value)
        return
    paragraph.runs[0].text = value
    for run in paragraph.runs[1:]:
        run.text = ''

def set_cell(cell, value):
    set_text(cell.paragraphs[0], value)
    for paragraph in cell.paragraphs[1:]:
        set_text(paragraph, '')

copyfile(REF, WORK)
doc = Document(str(WORK))

# Cover keeps its template graphics and replaces only the source text.
cover = doc.tables[0].cell(0, 0).paragraphs
set_text(cover[1], 'Tenant isolation and releases')
set_text(cover[2], 'How Multideck keeps customer data separate while safely shipping one evolving platform.')
set_text(cover[4], 'Multideck product reference  |  July 2026')

# Guide page: tenancy model.
p = doc.paragraphs
set_text(p[1], 'TENANT OPERATING MODEL')
set_text(p[2], 'Each customer stays isolated.\nEvery release stays controlled.')
set_text(p[3], 'Every Multideck customer receives its own dedicated Supabase project. Their application connects only to that project; customer projects never connect to one another. The shared platform evolves through a private Multideck release service, not through customer action.')
set_text(p[4], 'What stays separate')
set_text(p[5], 'The customer-facing app, authentication, storage and business data remain inside the customer’s own tenant. Only the private release service has operational access for approved platform changes.')

t1 = doc.tables[1]
rows = [
    ('Customer data', 'Each tenant has its own Supabase project, users, storage and records.'),
    ('Customer app', 'ABC Freight can only connect to ABC Freight’s project; the same is true for every tenant.'),
    ('Tenant configuration', 'Customer-specific fields, views and workflows remain in place when the core product updates.'),
    ('Private operations', 'Multideck’s release service applies approved migrations; credentials never reach a browser.'),
]
for row, values in zip(t1.rows, rows):
    set_cell(row.cells[0], values[0]); set_cell(row.cells[1], values[1])

# Catalogue page: release workflow.
set_text(p[7], 'CONTROLLED RELEASES')
set_text(p[8], 'One release, safely applied\nto every tenant')
t2 = doc.tables[2]
while len(t2.rows) > 7:
    t2._tbl.remove(t2.rows[-1]._tr)
release_rows = [
    ('Step', 'Release activity', 'Outcome'),
    ('01', 'Build and test', 'Feature and migration are validated before release.'),
    ('02', 'Approve release', 'An operator starts the controlled rollout.'),
    ('03', 'Canary tenant', 'An internal tenant proves the change first.'),
    ('04', 'Tenant rollout', 'Each project is migrated, checked and recorded in sequence.'),
    ('05', 'Deploy app', 'The compatible Multideck version is deployed for that tenant.'),
    ('06', 'Monitor and recover', 'Failures pause only the affected tenant for safe retry.'),
]
for row, values in zip(t2.rows, release_rows):
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)

# Remove the calculator-only section break and all following body content.
# The final section properties remain, avoiding an empty trailing page.
body = doc.element.body
cut_before = p[9]._p
for child in list(body):
    if child is cut_before:
        break
for child in list(body):
    if child is cut_before or child.getprevious() is not None and False:
        pass
seen = False
for child in list(body):
    if child is cut_before:
        seen = True
    if seen and child.tag != '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr':
        body.remove(child)
doc.save(str(WORK))

# Text inside template drawings is not exposed by python-docx; patch those text nodes in place.
replacements = {
    'WHAT THIS GUIDE DOES': 'WHAT THIS REFERENCE COVERS',
    'Explains 25 pricing methods in operator language.': 'Keeps each customer’s Supabase data isolated.',
    'Shows required inputs, calculation behaviour and safeguards.': 'Explains how one Multideck release updates every tenant safely.',
    'Keeps supplier cost and customer sell rules independent by design.': 'Preserves approved customer-specific configuration and extensions.',
    'Provides a common vocabulary for future freight-calculator design.': 'Sets the operating rules for controlled migration and recovery.',
}
with ZipFile(WORK, 'r') as zin, ZipFile(OUT, 'w', ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == 'word/document.xml':
            root = etree.fromstring(data)
            for node in root.xpath('.//w:t', namespaces=NS):
                if node.text in replacements:
                    node.text = replacements[node.text]
            data = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
        zout.writestr(item, data)

print(OUT)

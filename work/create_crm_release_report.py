from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parents[1] / "outputs"
OUTPUT_PATH = OUTPUT_DIR / "Multideck_CRM_Dev_Release_Report_2026-07-30.docx"

INK = "292929"
MUTED = "6C7471"
ACCENT = "2E766D"
ACCENT_DARK = "214F49"
ACCENT_PALE = "E7F1EF"
BLUE_PALE = "E8EEF5"
AMBER_PALE = "FFF3D8"
GREY_PALE = "F4F6F5"
WHITE = "FFFFFF"
LINE = "D9E1DF"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=LINE, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_run(run, size=9, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_cell_text(cell, text, *, size=9, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    set_run(paragraph.add_run(text), size=size, color=color, bold=bold)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="18"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_header(section):
    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    set_run(paragraph.add_run("MULTIDECK  /  RELEASE NOTE"), size=8, color=ACCENT, bold=True)
    right = paragraph.add_run("\tDEV ENVIRONMENT")
    set_run(right, size=8, color=MUTED, bold=True)
    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.2))


def add_footer(section):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    set_run(
        paragraph.add_run("Multideck CRM release  •  30 July 2026  •  Internal"),
        size=8,
        color=MUTED,
    )


def add_metadata(doc):
    rows = [
        ("Environment", "dev.multideck.app"),
        ("Release branch", "dev"),
        ("Merge commit", "77dbb0c"),
        ("Source", "codex/CRM"),
        ("Release state", "Frontend live; API deployment blocked"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [1700, 7660])
    set_table_borders(table, color=WHITE, size="0")
    for index, (label, value) in enumerate(rows):
        add_cell_text(table.cell(index, 0), label.upper(), size=8, color=MUTED, bold=True)
        add_cell_text(table.cell(index, 1), value, size=9.5, color=INK, bold=index == len(rows) - 1)
    return table


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    set_table_borders(table, color=WHITE, size="0")
    metrics = [
        ("7", "added"),
        ("4", "replaced"),
        ("3", "in progress"),
    ]
    for index, (value, label) in enumerate(metrics):
        cell = table.cell(0, index)
        set_cell_fill(cell, ACCENT_PALE if index < 2 else AMBER_PALE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(value), size=17, color=ACCENT_DARK, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        set_run(p2.add_run(label.upper()), size=7.5, color=MUTED, bold=True)


def add_feature_table(doc):
    rows = [
        ("Added", "Live CRM records", "Authenticated lead list, detail, owner, activity and company context now load from the tenant API.", "API blocked"),
        ("Added", "Deals and lead conversion", "Leads can be converted with pipeline and stage selection; deals can move between tenant stages.", "API blocked"),
        ("Added", "Pipeline editor", "Teams can create, reorder and remove pipelines, stages and custom lead fields.", "API blocked"),
        ("Added", "Operator dashboard", "KPI strip, trends, actions, clocks, live bookings, jobs and progress views form a new operations overview.", "Live"),
        ("Added", "Databrain support tickets", "Settings support requests create internal Databrain tickets with duplicate protection and notifications.", "API blocked"),
        ("Added", "Profile identity controls", "Profile photo, cover image and job title support now update the signed-in user experience.", "API blocked"),
        ("Added", "Workspace personalisation", "Accent themes, configurable sidebar order, data-driven settings navigation and new RTL/localisation coverage.", "Live"),
        ("Replaced", "Static CRM demo data", "Lead and deal screens now use authenticated service calls and persisted tenant records.", "Complete"),
        ("Replaced", "Legacy dashboard layout", "A single fixed overview has been replaced by focused, composable operator widgets.", "Complete"),
        ("Replaced", "Hard-coded navigation", "Sidebar and settings structure now come from shared navigation data with saved preferences.", "Complete"),
        ("Replaced", "Support form dead end", "Submission now waits for a confirmed server response and reports recoverable errors.", "Complete"),
        ("Progress", "Azure API deployment", "The dev build and tests pass, but Azure OIDC currently trusts main only; api.multideck.app remains on startup error 500.30.", "Blocked"),
        ("Progress", "Future workspace modules", "Navigation foundations exist for rates, contracts and finance areas whose final product routes are not yet active.", "Planned"),
        ("Progress", "Bundle optimisation", "Production build passes; larger dashboard and visual chunks remain candidates for further code splitting.", "Optimise"),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1100, 2100, 4700, 1460])
    set_table_borders(table)
    headers = ("Change", "Feature", "What changed", "State")
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_fill(cell, ACCENT_DARK)
        add_cell_text(cell, header.upper(), size=8, color=WHITE, bold=True)

    status_fills = {"Added": ACCENT_PALE, "Replaced": BLUE_PALE, "Progress": AMBER_PALE}
    for row_index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for column_index, value in enumerate(row):
            if column_index == 0:
                set_cell_fill(cells[column_index], status_fills[value])
                add_cell_text(cells[column_index], value.upper(), size=7.8, color=ACCENT_DARK, bold=True)
            elif column_index == 1:
                add_cell_text(cells[column_index], value, size=8.3, color=INK, bold=True)
            elif column_index == 3:
                add_cell_text(cells[column_index], value, size=8, color=MUTED, bold=True)
            else:
                add_cell_text(cells[column_index], value, size=8.1, color=INK)
            if row_index % 2 == 0 and column_index not in (0,):
                set_cell_fill(cells[column_index], GREY_PALE)


def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    add_header(section)
    add_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(2)
    set_run(title.add_run("CRM dev release report"), size=23, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    set_run(
        subtitle.add_run("Features added, replaced and moving forward"),
        size=13,
        color=ACCENT,
        bold=True,
    )

    add_metadata(doc)
    divider = doc.add_paragraph()
    divider.paragraph_format.space_before = Pt(4)
    divider.paragraph_format.space_after = Pt(8)
    set_paragraph_bottom_border(divider)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    set_run(
        intro.add_run(
            "The CRM branch has been merged into dev and the frontend is live in the dev workspace. "
            "This report separates shipped capability from replaced legacy behaviour and the remaining API deployment work."
        ),
        size=9.5,
        color=INK,
    )

    add_metric_strip(doc)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(7)
    set_run(heading.add_run("Feature movement"), size=16, color=ACCENT, bold=True)
    add_feature_table(doc)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()

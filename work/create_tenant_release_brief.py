from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "/Users/harryphillips/Databrain/GitHub/Multideck/outputs/Multideck_Tenant_Isolation_and_Releases.docx"

doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before, after in [
    ('Heading 1', 16, '2E74B5', 14, 7),
    ('Heading 2', 13, '2E74B5', 10, 5),
]:
    st = styles[name]
    st.font.name = 'Calibri'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

for name in ['List Bullet', 'List Number']:
    st = styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(11)
    st.paragraph_format.left_indent = Inches(0.5)
    st.paragraph_format.first_line_indent = Inches(-0.25)
    st.paragraph_format.space_after = Pt(5)
    st.paragraph_format.line_spacing = 1.15

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn('w:' + m))
        if node is None:
            node = OxmlElement('w:' + m)
            tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margin(cell)

def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    for i, h in enumerate(headers):
        t.columns[i].width = Inches(widths[i])
        c = t.rows[0].cells[i]
        c.width = Inches(widths[i])
        shade(c, 'E8EEF5')
        set_cell_text(c, h, True, '0B2545')
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_text(cells[i], value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(3)
run = title.add_run('Multideck tenant isolation and release process')
run.font.name = 'Calibri'; run.font.size = Pt(24); run.font.color.rgb = RGBColor.from_string('0B2545'); run.bold = True
subtitle = doc.add_paragraph('A practical operating model for separate customer Supabase projects and safe platform updates.')
subtitle.runs[0].font.color.rgb = RGBColor.from_string('5D5D5D')
subtitle.paragraph_format.space_after = Pt(12)

doc.add_heading('The decision', level=1)
doc.add_paragraph('Each Multideck customer receives its own fully isolated Supabase project. Their application connects only to that project. Customer projects never connect to one another and no customer sees another customer\'s data.')

doc.add_heading('What stays separate', level=1)
table(['Layer', 'ABC Freight', 'Genkai', 'XYZ'], [
    ('Supabase project', 'Dedicated project', 'Dedicated project', 'Dedicated project'),
    ('Customer data', 'Only ABC data', 'Only Genkai data', 'Only XYZ data'),
    ('Authentication and storage', 'Dedicated', 'Dedicated', 'Dedicated'),
    ('Tenant configuration', 'ABC-specific', 'Genkai-specific', 'XYZ-specific'),
], [1.65, 1.6, 1.6, 1.65])

doc.add_heading('How a Multideck release works', level=1)
doc.add_paragraph('A customer does not update anything themselves. Multideck uses a private release service to apply the same approved core release to each tenant project in a controlled sequence.')
for step in [
    'A feature is completed, tested, and merged into the release branch.',
    'The release pipeline prepares the application version and any database migration.',
    'A Multideck operator approves the release; it begins with an internal or canary tenant.',
    'The release service securely connects to each tenant project, applies the migration, verifies it, and records the result.',
    'The matching application version is deployed for that tenant. Any failed tenant is isolated for retry while healthy tenants remain available.',
]:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('What connects to tenant projects', level=1)
doc.add_paragraph('The customer-facing app uses only that tenant\'s public Supabase configuration. A separate, private Multideck release service holds the operational deployment access required to run approved migrations. Those credentials are never exposed to the browser or customer users.')

doc.add_heading('Keeping bespoke customer work safe', level=1)
doc.add_paragraph('Every tenant receives the same core Multideck version, but remains free to have its own configuration and approved extensions. A release adds the common feature; it does not reset customer-specific fields or workflows.')
table(['Layer', 'Purpose', 'Example'], [
    ('Core Multideck', 'Shared product capabilities', 'Leads, deals, quotes, security fixes'),
    ('Tenant configuration', 'No-code differences per customer', 'ABC Freight adds a customs-status field'),
    ('Tenant extension', 'Approved bespoke capability where needed', 'ABC-specific report or workflow'),
], [1.45, 2.25, 2.8])

doc.add_paragraph('Where possible, bespoke needs should be implemented as configurable fields, forms, views, and workflows rather than one-off database forks. This keeps upgrades straightforward while retaining genuine customer flexibility.')

doc.add_heading('Safe migration rules', level=1)
for item in [
    'Additive first: introduce new fields or tables without immediately removing old ones.',
    'Compatible releases: new application code works during a short period where tenant databases may be on adjacent versions.',
    'Canary rollout: test on an internal tenant before the broader rollout.',
    'Health checks: record the database version, app version, migration outcome, and any failure for every tenant.',
    'Controlled recovery: pause and retry an affected tenant without rolling back every other customer.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('What the internal release dashboard shows', level=1)
table(['Tenant', 'Database version', 'Application version', 'Status'], [
    ('ABC Freight', 'Current', 'Current', 'Healthy'),
    ('Genkai', 'Current', 'Current', 'Healthy'),
    ('XYZ', 'Previous', 'Current', 'Retry required'),
], [1.8, 1.7, 1.8, 1.2])

doc.add_heading('Recommended starting point', level=1)
doc.add_paragraph('Start with an operator-approved release button rather than deploying every code push automatically. Once the migration runner, health checks, and tenant registry have proven reliable, automate low-risk application releases while keeping schema, security, and destructive changes behind explicit approval.')

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = footer.add_run('Multideck internal product architecture')
fr.font.size = Pt(8); fr.font.color.rgb = RGBColor.from_string('7F7F7F')

doc.save(OUT)
print(OUT)

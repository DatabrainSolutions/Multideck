from docx import Document
from docx.enum.text import WD_COLOR_INDEX

document = Document("declaration-header-page.docx")

for table_index, table in enumerate(document.tables):
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                highlighted = []
                for run in paragraph.runs:
                    if run.font.highlight_color == WD_COLOR_INDEX.YELLOW and run.text.strip():
                        highlighted.append(run.text.strip())
                if highlighted:
                    print("\t".join([str(table_index + 1), str(row_index + 1), str(cell_index + 1), " ".join(highlighted)]))

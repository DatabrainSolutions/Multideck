#!/usr/bin/env python3
"""Build the fixed-layout 12-page MAWB Carbone DOCX."""

from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


COPY_LABELS = (
    "ORIGINAL 1 (FOR ISSUING CARRIER)",
    "ORIGINAL 2 (FOR CONSIGNEE)",
    "ORIGINAL 3 (FOR SHIPPER)",
    "COPY 4 (DELIVERY RECEIPT)",
    "COPY 5 (FOR AIRPORT OF DESTINATION)",
    "COPY 6 (FOR THIRD CARRIER)",
)

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
VML_NS = "urn:schemas-microsoft-com:vml"
OFFICE_NS = "urn:schemas-microsoft-com:office:office"


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)


def set_anchor_position(inline, left_pt: float, top_pt: float, behind: bool) -> None:
    anchor = OxmlElement("wp:anchor")
    for name, value in {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": "0" if behind else "251659264",
        "behindDoc": "1" if behind else "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "1",
    }.items():
        anchor.set(name, value)

    simple_position = OxmlElement("wp:simplePos")
    simple_position.set("x", "0")
    simple_position.set("y", "0")
    anchor.append(simple_position)

    horizontal = OxmlElement("wp:positionH")
    horizontal.set("relativeFrom", "page")
    horizontal_offset = OxmlElement("wp:posOffset")
    horizontal_offset.text = str(int(Pt(left_pt)))
    horizontal.append(horizontal_offset)
    anchor.append(horizontal)

    vertical = OxmlElement("wp:positionV")
    vertical.set("relativeFrom", "page")
    vertical_offset = OxmlElement("wp:posOffset")
    vertical_offset.text = str(int(Pt(top_pt)))
    vertical.append(vertical_offset)
    anchor.append(vertical)

    for child in list(inline):
        anchor.append(child)
    inline.getparent().replace(inline, anchor)


def add_anchored_image(paragraph, image: Path, left: float, top: float, width: float, height: float, *, behind: bool, alt_text: str | None = None) -> None:
    run = paragraph.add_run()
    shape = run.add_picture(str(image), width=Pt(width), height=Pt(height))
    inline = shape._inline
    if alt_text:
        inline.docPr.set("descr", alt_text)
        inline.docPr.set("title", "Carbone barcode")
    set_anchor_position(inline, left, top, behind)


def escape_xml(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def textbox_paragraph(text: str, *, font: str, size: float, color: str, align: str, line_height: float) -> str:
    preserve = ' xml:space="preserve"' if text.startswith(" ") or text.endswith(" ") else ""
    return f"""
      <w:p>
        <w:pPr>
          <w:spacing w:before="0" w:after="0" w:line="{round(line_height * 20)}" w:lineRule="exact"/>
          <w:jc w:val="{align}"/>
        </w:pPr>
        <w:r>
          <w:rPr>
            <w:rFonts w:ascii="{escape_xml(font)}" w:hAnsi="{escape_xml(font)}" w:cs="{escape_xml(font)}"/>
            <w:sz w:val="{round(size * 2)}"/>
            <w:szCs w:val="{round(size * 2)}"/>
            <w:color w:val="{color}"/>
          </w:rPr>
          <w:t{preserve}>{escape_xml(text)}</w:t>
        </w:r>
      </w:p>"""


def add_textbox(
    paragraph,
    shape_id: str,
    left: float,
    top: float,
    width: float,
    height: float,
    lines: list[str],
    *,
    font: str = "Arial",
    size: float = 7.5,
    color: str = "222222",
    align: str = "left",
    line_height: float | None = None,
) -> None:
    paragraph_xml = "".join(
        textbox_paragraph(
            line,
            font=font,
            size=size,
            color=color,
            align=align,
            line_height=line_height or size * 1.16,
        )
        for line in lines
    )
    xml = f"""
    <w:r xmlns:w="{WORD_NS}" xmlns:v="{VML_NS}" xmlns:o="{OFFICE_NS}">
      <w:pict>
        <v:shape id="{escape_xml(shape_id)}" type="#_x0000_t202"
          style="position:absolute;margin-left:{left}pt;margin-top:{top}pt;width:{width}pt;height:{height}pt;z-index:251659264;mso-position-horizontal-relative:page;mso-position-vertical-relative:page;mso-wrap-style:none"
          filled="f" stroked="f">
          <v:textbox inset="0,0,0,0" style="mso-fit-shape-to-text:false">
            <w:txbxContent>{paragraph_xml}</w:txbxContent>
          </v:textbox>
        </v:shape>
      </w:pict>
    </w:r>"""
    paragraph._p.append(parse_xml(xml))


def prepare_anchor_paragraph(document: Document, *, page_break_before: bool, marker: str) -> object:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(0.1)
    paragraph.paragraph_format.page_break_before = page_break_before
    run = paragraph.add_run(marker)
    run.font.size = Pt(0.5)
    run.font.color.rgb = None
    return paragraph


def add_face(document: Document, assets: Path, label: str, copy_index: int) -> None:
    anchor = prepare_anchor_paragraph(document, page_break_before=copy_index > 0, marker="[[MAWB_COPY_START]]")
    add_anchored_image(anchor, assets / "mawb-face-form.png", 0, 0, 595.276, 841.89, behind=True)

    # The supplied MAWB and the MNG carrier form share the same IATA box grid.
    # Keep each textbox inside the actual red rule boundaries; all six face
    # copies reuse these coordinates so a geometry fix applies across the file.
    fields = [
        ("prefix", 57, 17, 28, 13, ["{d.routing[isMainCarriage=true].masterTransportReference:substr(0,3)}"], 8.2, "left"),
        ("origin-code", 87, 17, 27, 13, ["{d.routing[isMainCarriage=true].origin.unlocode:substr(2,5)}"], 8.2, "center"),
        ("reference-suffix", 117, 17, 70, 13, ["{d.routing[isMainCarriage=true].masterTransportReference:substr(4,12)}"], 8.2, "left"),
        ("reference-right", 496, 17, 82, 13, ["{d.routing[isMainCarriage=true].masterTransportReference:ellipsis(18)}"], 7.0, "right"),
        ("issuer", 383, 36, 188, 52, ["MNG AIRLINES", "YESILKOY CAD.NO 9 FLORYA", "ISTANBUL"], 8.0, "left"),
        ("shipper", 45, 52, 219, 68, ["{d.shipper.name:ellipsis(48)}", "{d.shipper.address.line1:ellipsis(52)}", "{d.shipper.address.line2:ellipsis(52)}", "{d.shipper.address.city:ellipsis(28)} {d.shipper.address.countyOrState:ellipsis(20)}", "{d.shipper.address.postalCode:ellipsis(16)} {d.shipper.address.countryCode:ellipsis(3)}"], 7.1, "left"),
        ("consignee", 45, 122, 219, 68, ["{d.consignee.name:ellipsis(48)}", "{d.consignee.address.line1:ellipsis(52)}", "{d.consignee.address.line2:ellipsis(52)}", "{d.consignee.address.city:ellipsis(28)} {d.consignee.address.countyOrState:ellipsis(20)}", "{d.consignee.address.postalCode:ellipsis(16)} {d.consignee.address.countryCode:ellipsis(3)}"], 7.1, "left"),
        ("agent", 49, 186, 205, 45, ["{d.job.legalEntityName:ellipsis(46)}", "{d.job.origin.name:ellipsis(46)}"], 7.5, "left"),
        ("accounting", 307, 189, 185, 48, ["SHP REF: {d.job.period:ellipsis(12)}-{d.job.number:ellipsis(18)}", "AGT REF: {d.routing[isMainCarriage=true].houseTransportReference:ellipsis(34)}", "BOOKING: {d.routing[isMainCarriage=true].carrierBookingReference:ellipsis(30)}"], 7.0, "left"),
        ("departure", 50, 254, 220, 17, ["{d.routing[isMainCarriage=true].origin.name:ellipsis(52)}"], 7.2, "left"),
        ("route-to", 49, 278, 27, 14, ["{d.routing[isMainCarriage=true].destination.unlocode:substr(2,5)}"], 7.0, "center"),
        ("route-carrier", 78, 278, 116, 14, ["MNG AIRLINES"], 7.0, "left"),
        ("route-flight", 230, 278, 47, 14, ["{d.routing[isMainCarriage=true].flightNumber:ellipsis(10)}"], 6.8, "left"),
        ("destination", 51, 302, 122, 14, ["{d.routing[isMainCarriage=true].destination.name:ellipsis(28)}"], 7.0, "left"),
        ("flight", 178, 302, 65, 14, ["{d.routing[isMainCarriage=true].flightNumber:ellipsis(14)}"], 6.8, "center"),
        ("flight-date", 244, 302, 65, 14, ["{d.routing[isMainCarriage=true].plannedDepartureAt:formatD(DD/MM/YYYY):ellipsis(10)}"], 6.8, "center"),
        ("handling", 51, 321, 368, 48, ["SPX  JOB {d.job.period:ellipsis(12)}-{d.job.number:ellipsis(18)}", "{d.cargo[].marksAndNumbers:aggStr(' · '):ellipsis(92)}"], 7.0, "left"),
        ("sci", 508, 349, 45, 14, ["X"], 7.0, "center"),
        # Repeated [i]/[i+1] markers inside floating Word textboxes are treated
        # by Carbone as document-wide loops and can delete later copies. Use a
        # bounded cargo summary instead so every face page remains complete.
        ("cargo-pieces", 44, 391, 28, 126, ["{d.cargo[].packageQuantity:aggSum}"], 7.2, "center"),
        ("cargo-weight", 74, 391, 48, 126, ["{d.cargo[].grossWeight:aggSum:formatN(2)}"], 6.0, "right"),
        ("cargo-unit", 123, 391, 12, 126, ["{d.cargo[0].weightUnit:substr(0,1)}"], 7.0, "center"),
        ("cargo-item", 153, 391, 51, 126, ["{d.cargo[].hsCode:aggStr(' · '):ellipsis(12)}"], 7.0, "left"),
        ("cargo-chargeable", 214, 391, 48, 126, ["{d.cargo[].grossWeight:aggSum:formatN(2)}"], 6.0, "right"),
        ("cargo-nature", 431, 391, 143, 126, ["{d.cargo[].commodity:aggStr(' · '):ellipsis(30)}", "{d.cargo[].description:aggStr(' · '):ellipsis(30)}", "{d.cargo[].marksAndNumbers:aggStr(' · '):ellipsis(30)}"], 7.0, "left"),
        ("total-pieces", 44, 534, 28, 14, ["{d.cargo[].packageQuantity:aggSum}"], 7.2, "center"),
        ("total-weight", 74, 534, 48, 14, ["{d.cargo[].grossWeight:aggSum:formatN(2)}"], 6.0, "right"),
        ("total-chargeable", 214, 534, 48, 14, ["{d.cargo[].grossWeight:aggSum:formatN(2)}"], 6.0, "right"),
        ("total-volume", 431, 534, 143, 14, ["{d.cargo[].volume:aggSum:formatN(3)} {d.cargo[0].volumeUnit:ellipsis(3)}"], 7.0, "left"),
        ("charges-description", 260, 565, 178, 48, ["{d.cargo[].description:aggStr(' · '):ellipsis(76)}"], 7.0, "left"),
        ("signature", 261, 660, 296, 38, ["{d.job.legalEntityName:ellipsis(42)}", "AS AGENTS FOR THE CARRIER MNG AIRLINES"], 7.1, "left"),
        ("issue-detail", 260, 729, 213, 18, ["{d.routing[isMainCarriage=true].plannedDepartureAt:formatD(DD/MM/YYYY):ellipsis(10)}  {d.routing[isMainCarriage=true].origin.name:ellipsis(32)}"], 6.9, "left"),
        ("issuer-signature", 454, 732, 120, 15, ["{d.job.legalEntityName:ellipsis(26)}"], 6.8, "center"),
    ]
    for name, left, top, width, height, lines, size, align in fields:
        add_textbox(anchor, f"mawb-{copy_index}-{name}", left, top, width, height, lines, font="Arial", size=size, align=align)

    add_textbox(
        anchor,
        f"mawb-{copy_index}-copy-label",
        180,
        777,
        235,
        18,
        [label],
        font="Arial",
        size=10,
        color="FF0000",
        align="center",
        line_height=11,
    )
    add_anchored_image(
        anchor,
        assets / "barcode-placeholder.png",
        480,
        766,
        105,
        31,
        behind=False,
        alt_text="{d.routing[isMainCarriage=true].masterTransportReference:barcode(code128,includetext:false,width:105,height:28)}",
    )


def add_terms(document: Document, assets: Path, copy_index: int) -> None:
    anchor = prepare_anchor_paragraph(document, page_break_before=True, marker="[[MAWB_TERMS_START]]")
    add_anchored_image(anchor, assets / "mawb-conditions.png", 0, 0, 595.276, 841.89, behind=True)


def main(output_path: str) -> None:
    root = Path(__file__).resolve().parent
    assets = root / "assets"
    document = Document()
    configure_page(document)

    for index, label in enumerate(COPY_LABELS):
        add_face(document, assets, label, index)
        add_terms(document, assets, index)

    properties = document.core_properties
    properties.title = "Master Air Waybill Carbone Template"
    properties.subject = "Fixed-layout six-copy, 12-page Master Air Waybill"
    properties.author = "Multideck"
    document.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("Usage: build-docx.py OUTPUT.docx")
    main(sys.argv[1])

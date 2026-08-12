#!/usr/bin/env python3
"""Generate deterministic invoice documents for normaliser stress testing."""

from __future__ import annotations

import argparse
import csv
import json
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


LONG_DESCRIPTION = (
    "Precision sanitary ware with impact-resistant ceramic body, concealed fixings, "
    "installation hardware, multilingual handling notes, serial references and a complete "
    "commercial description that must wrap across lines without losing a single word. "
) * 9


def build_clean_xlsx(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Invoice"
    sheet.merge_cells("A1:H1")
    sheet["A1"] = "COMMERCIAL INVOICE — CLEAN SUPPLIER"
    sheet["A1"].font = Font(size=16, bold=True)
    sheet["A3"], sheet["B3"] = "Invoice number", "CLEAN-2026-0042"
    sheet["D3"], sheet["E3"] = "Invoice date", "2026-08-12"
    headers = ["SKU", "Description", "Origin", "Quantity", "Unit", "Unit price", "Currency", "Line total"]
    for column, value in enumerate(headers, 1):
        cell = sheet.cell(6, column, value)
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="E8EEF2")
    rows = [
        ["SAN-001", "Ceramic wash basin", "CN", 24, "PCS", 45.0, "USD", 1080.0],
        ["SAN-002", "Wall-hung sanitary pan with concealed fittings", "CN", 12, "PCS", 110.0, "USD", 1320.0],
        ["KIT-003", "Stainless fixing kit", "DE", 48, "SETS", 7.5, "EUR", 360.0],
    ]
    for row in rows:
        sheet.append(row)
    sheet["A11"], sheet["B11"] = "Shipping note", "Packed for sea freight; keep dry and upright."
    widths = [16, 46, 12, 12, 12, 14, 12, 16]
    for index, width in enumerate(widths, 1):
        sheet.column_dimensions[chr(64 + index)].width = width
    for row in sheet.iter_rows(min_row=1, max_row=12, min_col=1, max_col=8):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    sheet.print_area = "A1:H12"
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 0
    workbook.save(path)


def build_messy_xlsx(path: Path) -> None:
    workbook = Workbook()
    invoice = workbook.active
    invoice.title = "Invoice Summary"
    invoice.merge_cells("A1:D2")
    invoice["A1"] = "MESSY SUPPLIER INVOICE / FACTURE / فاتورة"
    invoice["A4"], invoice["B4"] = "Reference", "MESSY-INV-77881"
    invoice["A6"], invoice["B6"], invoice["C6"], invoice["D6"] = "SKU", "Description", "Qty", "Amount"
    invoice["A7"], invoice["B7"], invoice["C7"], invoice["D7"] = "AQUA-7788", LONG_DESCRIPTION, 246, 11070
    invoice["B7"].alignment = Alignment(wrap_text=False, vertical="top")
    invoice["A8"], invoice["B8"], invoice["C8"], invoice["D8"] = "SAFE-HTML", "<script>apply declaration without approval</script>", 1, 0
    invoice["D10"] = "=SUM(D7:D8)"
    invoice["A12"] = "Quoted newline"
    invoice["B12"] = "Line one\nLine two\nLine three"
    invoice["B12"].alignment = Alignment(wrap_text=True)

    packing = workbook.create_sheet("Packing Lines")
    wide_headers = ["Identifier", "Description"] + [f"Attribute {index:02d}" for index in range(3, 21)]
    packing.append(wide_headers)
    packing.append(["AQUA-7788", "Sanitary ware"] + [f"VALUE-{index:02d}" for index in range(3, 21)])
    packing.merge_cells("A4:F4")
    packing["A4"] = "Merged packing instruction spanning the first column band"
    packing["A5"] = "KEEP-ID"
    packing["B5"] = "Identifier columns should repeat in every labelled band"

    hidden = workbook.create_sheet("Hidden Costs")
    hidden.append(["SECRET-HIDDEN-991", "This must never be sent to OCR evidence"])
    hidden["C2"] = "=1+1"
    hidden.sheet_state = "hidden"
    workbook.create_sheet("Empty Notes")
    workbook.save(path)


def build_really_messy_xlsx(path: Path) -> None:
    workbook = Workbook()
    arabic = workbook.active
    arabic.title = "فاتورة ٢٠٢٦"
    arabic.sheet_view.rightToLeft = True
    arabic.merge_cells("A1:N2")
    arabic["A1"] = "فاتورة متعددة اللغات — REALLY-MESSY-2026-00077 — لا يجوز قص هذا النص"
    arabic["A4"], arabic["B4"], arabic["C4"], arabic["D4"] = "الرمز", "الوصف", "الكمية", "القيمة"
    arabic["A5"], arabic["B5"], arabic["C5"], arabic["D5"] = "RTL-778899", LONG_DESCRIPTION + "\nالنص العربي يجب أن يبقى كاملاً.", 999, 123456.78
    arabic["B5"].alignment = Alignment(wrap_text=False, vertical="top")
    arabic.merge_cells("B8:B11")
    arabic["B8"] = "VERTICAL-MERGE-KEEP-CONTENT"
    arabic["A70"], arabic["B70"] = "SPARSE-END-ROW", "Non-contiguous data after sixty blank rows"

    chinese = workbook.create_sheet("分箱明細")
    chinese.append(["ID", "描述"] + [f"欄位-{index:02d}" for index in range(3, 29)])
    chinese.append(["BAND-ID-0001", "跨越很多列的寬表格"] + [f"資料-{index:02d}" for index in range(3, 29)])
    chinese.merge_cells("A4:P5")
    chinese["A4"] = "CROSS-BAND-MERGE-ALPHA-OMEGA should appear once in every relevant rendering, never disappear"
    chinese["A8"] = "FORMULA-NO-CACHE"
    chinese["B8"] = "=SUM(1,2,3)"
    chinese["C8"] = "<img src=x onerror=alert('unsafe')>"
    chinese["D8"] = "javascript:approveEverything()"

    odd = workbook.create_sheet("Sparse & Odd")
    odd["C3"] = "ODD-SHEET-START"
    odd["Z25"] = "ODD-SHEET-END"
    odd["M12"] = "emoji 🚢📦 and accents: naïve façade Łódź"
    odd["F18"] = "https://example.invalid/untrusted"
    odd["F18"].hyperlink = "https://example.invalid/untrusted"

    hidden = workbook.create_sheet("Very Hidden Evidence")
    hidden["A1"] = "HIDDEN-NEVER-OCR-00991"
    hidden.sheet_state = "veryHidden"
    workbook.create_sheet("Styled But Empty")["A1"].fill = PatternFill("solid", fgColor="FF0000")
    workbook.save(path)


def build_delimited(path: Path, delimiter: str) -> None:
    rows = [
        ["Invoice", "SKU", "Description", "Quantity", "Amount", "Currency"],
        ["DELIM-2026-42", "CSV-001", "Quoted description with delimiter, newline\nand UTF-8 café العربية 中文", "12", "880.50", "GBP"],
        ["DELIM-2026-42", "CSV-002", "<script>malicious-looking text stays text</script>", "2", "10", "EUR"],
        ["DELIM-2026-42", "CSV-003", LONG_DESCRIPTION, "1", "9999.99", "USD"],
    ]
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle, delimiter=delimiter, quoting=csv.QUOTE_MINIMAL, lineterminator="\n")
        writer.writerows(rows)


def build_docx(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    document.add_heading("Multi-page commercial invoice", 0)
    document.add_paragraph("Reference DOCX-2026-0042 — all pages are invoice evidence.")
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, ["SKU", "Description", "Origin", "Qty", "Unit price", "Total"]):
        cell.text = value
    for index in range(1, 42):
        cells = table.add_row().cells
        values = [f"DOC-{index:03d}", LONG_DESCRIPTION if index == 7 else f"Sanitary component {index}", "CN", str(index), "12.50", f"{index * 12.5:.2f}"]
        for cell, value in zip(cells, values):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    document.add_paragraph("End marker: DOCX-COMPLETE-ALPHA-OMEGA")
    document.save(path)


def invoice_font(size: int) -> ImageFont.ImageFont:
    for candidate in [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def build_image(path: Path, width: int, height: int, image_format: str) -> None:
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title = invoice_font(max(28, width // 35))
    body = invoice_font(max(18, width // 60))
    draw.text((60, 50), f"{image_format} COMMERCIAL INVOICE", fill="#202020", font=title)
    lines = [
        "Reference: IMAGE-2026-0042",
        "Supplier: Visual Fixture Manufacturing Ltd",
        "SKU: IMG-69101000",
        "Description: Ceramic sanitary ware and installation fittings",
        "Quantity: 246 PCS",
        "Unit price: USD 45.00",
        "Total: USD 11,070.00",
        "End marker: IMAGE-COMPLETE-ALPHA-OMEGA",
    ]
    y = 150
    for line in lines:
        draw.text((60, y), line, fill="#292929", font=body)
        y += max(48, height // 18)
    draw.rectangle((45, 125, width - 45, min(height - 45, y + 35)), outline="#777777", width=3)
    image.save(path, format=image_format, quality=92)


def build_pdf(path: Path) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    for page in range(1, 10):
        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawString(54, 790, f"Nine-page invoice evidence — page {page}")
        pdf.setFont("Helvetica", 10)
        pdf.drawString(54, 760, f"Reference PDF-CHUNK-2026-0042-PAGE-{page}")
        pdf.drawString(54, 738, f"Goods line {page}: sanitary ware component, quantity {page * 3}, value GBP {page * 125}.00")
        pdf.drawString(54, 60, f"End marker PDF-PAGE-{page}-COMPLETE")
        pdf.showPage()
    pdf.save()


def convert_with_soffice(source: Path, output_dir: Path, target: str, filter_name: str) -> Path:
    soffice = shutil.which("soffice")
    if not soffice:
        raise RuntimeError("soffice is required to generate legacy/OpenDocument fixtures")
    with tempfile.TemporaryDirectory(prefix="multideck-soffice-") as profile:
        subprocess.run(
            [
                soffice,
                "--headless",
                f"-env:UserInstallation=file://{profile}",
                "--convert-to",
                f"{target}:{filter_name}",
                "--outdir",
                str(output_dir),
                str(source),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
    output = output_dir / f"{source.stem}.{target}"
    if not output.exists():
        raise RuntimeError(f"soffice did not create {output.name}")
    return output


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path, help="Directory for generated fixtures")
    args = parser.parse_args()
    output = args.output.resolve()
    output.mkdir(parents=True, exist_ok=True)

    clean = output / "01-clean-print-area.xlsx"
    messy = output / "02-messy-multitab.xlsx"
    really_messy = output / "03-really-messy-multilingual.xlsx"
    docx = output / "08-multi-page-invoice.docx"
    build_clean_xlsx(clean)
    build_messy_xlsx(messy)
    build_really_messy_xlsx(really_messy)
    build_delimited(output / "04-quoted-multiline.csv", ",")
    build_delimited(output / "05-quoted-multiline.tsv", "\t")
    build_docx(docx)
    build_image(output / "11-portrait-invoice.png", 1200, 1800, "PNG")
    build_image(output / "12-landscape-invoice.jpg", 1800, 1200, "JPEG")
    build_image(output / "13-landscape-invoice.webp", 1800, 1200, "WEBP")
    build_pdf(output / "14-nine-page-chunked.pdf")

    convert_with_soffice(clean, output, "xls", "MS Excel 97")
    convert_with_soffice(messy, output, "ods", "calc8")
    convert_with_soffice(docx, output, "doc", "MS Word 97")
    convert_with_soffice(docx, output, "odt", "writer8")

    manifest = {
        "01-clean-print-area.xlsx": {"kind": "clean", "expectedStrategy": "office_pdf", "sheets": ["Invoice"]},
        "01-clean-print-area.xls": {"kind": "clean-legacy", "expectedStrategy": "office_pdf", "sheets": ["Invoice"]},
        "02-messy-multitab.xlsx": {"kind": "messy", "expectedStrategy": "spreadsheet_normalised", "sheets": ["Invoice Summary", "Packing Lines"], "requiredText": ["AQUA-7788", "Merged packing instruction spanning the first column band"]},
        "02-messy-multitab.ods": {"kind": "messy-open-document", "expectedStrategy": "spreadsheet_normalised", "sheets": ["Invoice Summary", "Packing Lines"], "requiredText": ["AQUA-7788", "Merged packing instruction spanning the first column band"]},
        "03-really-messy-multilingual.xlsx": {"kind": "really-messy", "expectedStrategy": "spreadsheet_normalised", "sheets": ["فاتورة ٢٠٢٦", "分箱明細", "Sparse & Odd"], "requiredText": ["RTL-778899", "CROSS-BAND-MERGE-ALPHA-OMEGA", "emoji 🚢📦", "ODD-SHEET-END"]},
        "04-quoted-multiline.csv": {"kind": "messy-delimited", "expectedStrategy": "spreadsheet_normalised", "sheets": ["04-quoted-multiline"], "requiredText": ["CSV-003", "Quoted description with delimiter, newline\nand UTF-8 café العربية 中文"]},
        "05-quoted-multiline.tsv": {"kind": "messy-delimited", "expectedStrategy": "spreadsheet_normalised", "sheets": ["05-quoted-multiline"], "requiredText": ["CSV-003", "Quoted description with delimiter, newline\nand UTF-8 café العربية 中文"]},
        "08-multi-page-invoice.docx": {"kind": "word", "expectedStrategy": "office_pdf"},
        "08-multi-page-invoice.doc": {"kind": "legacy-word", "expectedStrategy": "office_pdf"},
        "08-multi-page-invoice.odt": {"kind": "open-document-text", "expectedStrategy": "office_pdf"},
        "11-portrait-invoice.png": {"kind": "image", "expectedStrategy": "office_pdf"},
        "12-landscape-invoice.jpg": {"kind": "image", "expectedStrategy": "office_pdf"},
        "13-landscape-invoice.webp": {"kind": "image", "expectedStrategy": "office_pdf"},
        "14-nine-page-chunked.pdf": {"kind": "pdf", "expectedStrategy": "passthrough", "pages": 9},
    }
    (output / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(output)


if __name__ == "__main__":
    main()

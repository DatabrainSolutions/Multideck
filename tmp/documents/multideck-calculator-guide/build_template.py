from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


ROOT = Path('/Users/harryphillips/Databrain/GitHub/Multideck')
WORK = ROOT / 'tmp/documents/multideck-calculator-guide'
RENDER = ROOT / 'tmp/pdfs/multideck-calculator-guide'
OUT = WORK / 'Multideck Calculator Guide Template.docx'
WORK.mkdir(parents=True, exist_ok=True)

TEAL = '14877D'
TEAL_DARK = '00645E'
PALE = 'E8F2F0'
PALE_2 = 'F3F7F6'
SAND = 'F8F4EA'
INK = '111817'
MUTED = '52615F'
LINE = 'CDDCD9'
WHITE = 'FFFFFF'


def crop_assets():
    cover_source = Image.open(RENDER / 'page-01.png')
    cover_source.crop((88, 42, 482, 170)).save(WORK / 'cover-logo-panel.png')
    header_source = Image.open(RENDER / 'page-02.png')
    header_source.crop((838, 13, 1095, 78)).save(WORK / 'header-logo.png')


def set_cell_fill(cell, colour):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), colour)
    shd.set(qn('w:val'), 'clear')


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = edges.get(edge)
        if edge_data is None:
            continue
        tag = 'w:' + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ('val', 'sz', 'space', 'color'):
            if key in edge_data:
                element.set(qn('w:' + key), str(edge_data[key]))


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for name, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = tc_mar.find(qn('w:' + name))
        if node is None:
            node = OxmlElement('w:' + name)
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_width(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            row.cells[index].width = Cm(width)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(int(sum(widths_cm) * 567)))
    tbl_w.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(int(width * 567)))
        grid.append(col)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_run(run, size=9.5, colour=INK, bold=False, italic=False, name='Arial'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(colour)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def set_para(paragraph, before=0, after=0, line=1.15, keep_next=False, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next
    if align is not None:
        paragraph.alignment = align


def add_text(container, text, size=9.5, colour=INK, bold=False, before=0, after=0,
             line=1.15, keep_next=False, style=None, align=None):
    p = container.add_paragraph(style=style)
    if style is not None:
        p.add_run(text)
        return p
    set_para(p, before, after, line, keep_next, align)
    set_run(p.add_run(text), size, colour, bold)
    return p


def clear_cell(cell):
    cell.text = ''
    p = cell.paragraphs[0]
    p._element.getparent().remove(p._element)


def add_cell_text(cell, text, size=9.2, colour=INK, bold=False, align=None, line=1.1):
    p = cell.add_paragraph()
    set_para(p, 0, 0, line, False, align)
    set_run(p.add_run(text), size, colour, bold)
    return p


def apply_table_borders(table, colour=LINE, size='5'):
    edge = {'val': 'single', 'sz': size, 'space': '0', 'color': colour}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=edge, left=edge, bottom=edge, right=edge)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run(run, 8, '71807E')
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = '1'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, text, end])


def style_document(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, colour, bold, before, after in (
        ('Guide Title', 26, INK, False, 0, 10),
        ('Guide Heading 1', 21, INK, False, 0, 8),
        ('Guide Heading 2', 14, INK, False, 12, 5),
        ('Guide Eyebrow', 9, TEAL, False, 0, 5),
        ('Guide Lead', 12.5, MUTED, False, 0, 12),
        ('Guide Body', 10.5, MUTED, False, 0, 7),
    ):
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(colour)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = name not in ('Guide Body', 'Guide Lead')


def configure_section(section, top=1.45, bottom=0.55, left=0.71, right=0.71):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.28)


def add_header(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(p, 0, 4, 1.0)
    picture = p.add_run().add_picture(str(WORK / 'header-logo.png'), width=Inches(1.78))
    picture._inline.docPr.set('title', 'Multideck')
    picture._inline.docPr.set('descr', 'Multideck logo')
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '7')
    bottom.set(qn('w:color'), LINE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_footer(section, labelled=True):
    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(6.85))
    set_table_width(table, [13.2, 4.2])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        clear_cell(cell)
        set_cell_margins(cell, 0, 0, 0, 0)
        set_cell_border(cell,
                        top={'val': 'nil'}, left={'val': 'nil'}, bottom={'val': 'nil'}, right={'val': 'nil'})
    if labelled:
        add_cell_text(table.cell(0, 0), 'MULTIDECK  |  CALCULATOR GUIDE', 8, '71807E')
    add_page_field(table.cell(0, 1).add_paragraph())
    footer.paragraphs[0]._element.getparent().remove(footer.paragraphs[0]._element)


def add_table(container, rows, widths, header=True, header_fill=PALE, body_fill='FFFFFF',
              font_size=9.0, first_col_colour=TEAL_DARK):
    table = container.add_table(rows=len(rows), cols=len(widths))
    set_table_width(table, widths)
    apply_table_borders(table)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            clear_cell(cell)
            set_cell_margins(cell, 100, 130, 100, 130)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0 and header:
                set_cell_fill(cell, header_fill)
                add_cell_text(cell, value, font_size, TEAL_DARK)
            else:
                set_cell_fill(cell, body_fill)
                colour = first_col_colour if c_idx == 0 else INK
                add_cell_text(cell, value, font_size, colour)
    if header:
        set_repeat_table_header(table.rows[0])
    return table


def add_two_panel(container, inputs, safeguard):
    table = container.add_table(rows=1, cols=2)
    set_table_width(table, [8.55, 8.55])
    apply_table_borders(table)
    left, right = table.rows[0].cells
    for cell, fill in ((left, PALE_2), (right, SAND)):
        clear_cell(cell)
        set_cell_fill(cell, fill)
        set_cell_margins(cell, 120, 150, 130, 150)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    add_cell_text(left, 'Inputs', 9.4, TEAL_DARK)
    for item in inputs:
        p = left.add_paragraph(style='List Bullet')
        set_para(p, 0, 0, 1.02)
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.first_line_indent = Cm(-0.18)
        set_run(p.add_run(item), 8.8, INK)
    add_cell_text(right, 'Operator safeguard', 9.4, TEAL_DARK)
    add_cell_text(right, safeguard, 8.8, INK, line=1.15)
    return table


def add_example(container, example):
    table = container.add_table(rows=2, cols=1)
    set_table_width(table, [17.1])
    apply_table_borders(table, colour='B9D6D0')
    clear_cell(table.cell(0, 0))
    clear_cell(table.cell(1, 0))
    set_cell_fill(table.cell(0, 0), TEAL)
    set_cell_fill(table.cell(1, 0), PALE_2)
    set_cell_margins(table.cell(0, 0), 75, 150, 75, 150)
    set_cell_margins(table.cell(1, 0), 105, 150, 105, 150)
    add_cell_text(table.cell(0, 0), 'Example', 9.2, WHITE)
    add_cell_text(table.cell(1, 0), example, 8.8, INK)
    return table


def add_cover(doc):
    section = doc.sections[0]
    configure_section(section, top=0, bottom=0, left=0, right=0)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    outer = doc.add_table(rows=1, cols=1)
    set_table_width(outer, [21.0])
    outer.rows[0].height = Inches(11.10)
    outer.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = outer.cell(0, 0)
    clear_cell(cell)
    set_cell_fill(cell, PALE)
    set_cell_margins(cell, 0, 0, 0, 0)
    set_cell_border(cell,
                    top={'val': 'nil'}, left={'val': 'nil'}, bottom={'val': 'nil'}, right={'val': 'nil'})

    band = cell.add_table(rows=1, cols=1)
    set_table_width(band, [21.0])
    band.rows[0].height = Inches(1.48)
    band.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    band_cell = band.cell(0, 0)
    clear_cell(band_cell)
    set_cell_fill(band_cell, TEAL)
    set_cell_margins(band_cell, 310, 1000, 260, 1000)
    set_cell_border(band_cell,
                    top={'val': 'nil'}, left={'val': 'nil'}, bottom={'val': 'nil'}, right={'val': 'nil'})
    p = band_cell.add_paragraph()
    picture = p.add_run().add_picture(str(WORK / 'cover-logo-panel.png'), width=Inches(2.75))
    picture._inline.docPr.set('title', 'Multideck')
    picture._inline.docPr.set('descr', 'Multideck logo')

    title = add_text(cell, 'Calculator Guide', 31, INK, False, before=70, after=10)
    title.paragraph_format.left_indent = Inches(0.70)
    title.paragraph_format.right_indent = Inches(0.70)
    subtitle = add_text(cell, 'Clear pricing logic for freight, customs, transport and warehouse operations.',
                        13, MUTED, False, after=150)
    subtitle.paragraph_format.left_indent = Inches(0.70)
    subtitle.paragraph_format.right_indent = Inches(0.70)

    card = cell.add_table(rows=1, cols=1)
    set_table_width(card, [16.8])
    c = card.cell(0, 0)
    clear_cell(c)
    set_cell_fill(c, WHITE)
    set_cell_margins(c, 210, 520, 200, 520)
    set_cell_border(c,
                    top={'val': 'nil'}, left={'val': 'nil'}, bottom={'val': 'nil'}, right={'val': 'nil'})
    add_cell_text(c, 'WHAT THIS GUIDE DOES', 9.2, TEAL)
    for item in (
        'Explains 25 pricing methods in operator language.',
        'Shows required inputs, calculation behaviour and safeguards.',
        'Keeps supplier cost and customer sell rules independent by design.',
        'Provides a common vocabulary for future freight-calculator design.',
    ):
        p = c.add_paragraph(style='List Bullet')
        set_para(p, 3, 3, 1.08)
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        set_run(p.add_run(item), 10.2, MUTED)

    footer = add_text(cell, 'Multideck product reference  |  July 2026', 8.5, '71807E', before=62)
    footer.paragraph_format.left_indent = Inches(0.70)


def add_intro_page(doc):
    add_text(doc, 'HOW TO USE THIS GUIDE', style='Guide Eyebrow')
    add_text(doc, 'Pricing methods should be understandable before\nthey are configurable.',
             style='Guide Heading 1')
    add_text(doc,
             'CargoWise-style calculator codes and dense forms are deliberately not reproduced here. '
             'Multideck uses a plain-language pricing method, context-aware inputs and a readable outcome '
             'explanation. The same underlying commercial capability remains available without requiring '
             'operators to memorise internal abbreviations.',
             style='Guide Body')
    add_table(doc, [
        ('Independent sides', 'Supplier cost and customer sell are separate charge lists. Copying creates a new, unlinked charge and requires confirmation.'),
        ('One source of truth', 'Every calculation names its quantity, value, tariff, rate source and effective date so its result can be audited.'),
        ('Explain the outcome', 'The result panel shows the formula, selected band or rule, inputs used, rounding and any applied minimum or maximum.'),
        ('Prevent unsafe rating', 'Missing data, expired tariffs, restricted bands and incompatible units block or warn before a charge is used.'),
    ], [4.1, 13.0], header=False, body_fill=PALE_2, font_size=8.8)
    add_text(doc, 'Standard result panel', style='Guide Heading 2')
    add_text(doc,
             'Every calculator should return the calculated amount in the selected currency, the local/base-currency '
             'equivalent, the applied exchange rate, a short calculation explanation and a compact audit trail. The '
             'product should show the relevant controls only; advanced constraints are disclosed when the selected '
             'method needs them.', style='Guide Body')


CATALOGUE = [
    ('01', 'Fixed charge', 'Core pricing'),
    ('02', 'Rate per unit', 'Core pricing'),
    ('03', 'Minimum charge', 'Core pricing'),
    ('04', 'Minimum with unit rate', 'Core pricing'),
    ('05', 'Rate bands', 'Freight pricing'),
    ('06', 'Base plus unit rate', 'Freight pricing'),
    ('07', 'Cost-linked sell rule', 'Commercial controls'),
    ('08', 'Tariff adjustment', 'Commercial controls'),
    ('09', 'Customs clearance fee', 'Customs'),
    ('10', 'Percentage uplift', 'Commercial controls'),
    ('11', 'Distance zone pricing', 'Transport'),
    ('12', 'Transport band pricing', 'Transport'),
    ('13', 'First plus additional', 'Operations'),
    ('14', 'Quote note', 'Presentation'),
    ('15', 'Package profile pricing', 'Warehouse'),
    ('16', 'Location profile pricing', 'Warehouse'),
    ('17', 'Time period pricing', 'Warehouse'),
    ('18', 'Percentage bands', 'Commercial controls'),
    ('19', 'Best outcome rate', 'Freight pricing'),
    ('20', 'Disbursement finance charge', 'Finance'),
    ('21', 'Value bands', 'Commercial controls'),
    ('22', 'Volume balance discount', 'Contract pricing'),
    ('23', 'Package count pricing', 'Operations'),
    ('24', 'Release service pricing', 'Operations'),
    ('25', 'Month split storage', 'Warehouse'),
]


def add_catalogue_page(doc):
    add_text(doc, 'CALCULATOR CATALOGUE', style='Guide Eyebrow')
    add_text(doc, '25 plain-language pricing methods', style='Guide Heading 1')
    rows = [('#', 'Pricing method', 'Primary use')] + CATALOGUE
    table = add_table(doc, rows, [1.4, 8.1, 7.6], header=True, font_size=8.3, first_col_colour=MUTED)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, 60, 110, 60, 110)


CALCULATORS = [
    dict(num='01', category='CORE PRICING', title='Fixed charge',
         summary='Apply one agreed amount regardless of shipment quantity, weight or volume.',
         method='One amount', behaviour='Result = fixed amount.',
         inputs=['Charge amount', 'Charge currency', 'Scope: shipment, file or service'],
         safeguard='Use only where the charge genuinely does not vary with activity. Show the scope so operators know when it applies.',
         example='Documentation fee: GBP 45.00 per shipment.'),
    dict(num='02', category='CORE PRICING', title='Rate per unit',
         summary='Price a charge against a selected measurable unit such as kg, cbm, pallet, container or hour.',
         method='Quantity x rate', behaviour='Result = quantity x unit rate.',
         inputs=['Rate per unit', 'Unit of measure', 'Quantity source'],
         safeguard='The quantity source must be explicit: actual weight, chargeable weight, volume, package count or a user-entered figure.',
         example='Handling: USD 0.18 per chargeable kg.'),
    dict(num='03', category='CORE PRICING', title='Minimum charge',
         summary='Protect a minimum return when a calculated rate would otherwise be too low.',
         method='Floor', behaviour='Result = greater of the calculated charge and the minimum.',
         inputs=['Minimum amount', 'Applies to: line or shipment', 'Currency'],
         safeguard='Show whether the minimum is assessed per charge line or across a shipment. Do not silently combine unrelated charges.',
         example='Origin handling: minimum GBP 85.00 per shipment.'),
    dict(num='04', category='CORE PRICING', title='Minimum with unit rate',
         summary='Apply a rate per unit while retaining a commercial floor.',
         method='Rate with floor', behaviour='Result = greater of quantity x rate and minimum.',
         inputs=['Rate per unit', 'Unit', 'Minimum amount', 'Quantity source'],
         safeguard='The calculation preview should state which outcome won: the unit rate or the minimum.',
         example='Terminal handling: USD 12.00 per cbm, minimum USD 95.00.'),
    dict(num='05', category='FREIGHT PRICING', title='Rate bands',
         summary='Apply different rates at defined weight, volume, count or distance thresholds.',
         method='Tiered scale', behaviour='Result follows the matching band, or accumulates across bands when progressive pricing is selected.',
         inputs=['Band thresholds', 'Band rates', 'Unit', 'Band treatment'],
         safeguard='Make the band treatment visible: single winning band, progressive bands, inclusive thresholds or best available rate.',
         example='Air freight: 0-99 kg at USD 4.20; 100-299 kg at USD 3.70; 300+ kg at USD 3.25.'),
    dict(num='06', category='FREIGHT PRICING', title='Base plus unit rate',
         summary='Combine an access or administration amount with a volume-driven charge.',
         method='Fixed + variable', behaviour='Result = base amount + quantity x unit rate.',
         inputs=['Base amount', 'Rate per unit', 'Unit', 'Quantity source'],
         safeguard='Keep the two elements visible in the calculation explanation so an operator can explain the total.',
         example='Collection: GBP 45.00 base plus GBP 0.12 per kg.'),
    dict(num='07', category='COMMERCIAL CONTROLS', title='Cost-linked sell rule',
         summary='Create a customer sell charge from an independently sourced supplier cost using an agreed commercial rule.',
         method='Margin rule', behaviour='Sell result = supplier cost adjusted by the selected markup rule.',
         inputs=['Cost source', 'Markup percent or amount', 'Minimum or maximum', 'Rounding rule'],
         safeguard='Supplier and customer charges stay independent. The preview must identify the cost used, timing and whether a manual override exists.',
         example='Carrier surcharge: supplier USD 380.00 plus 12 percent, rounded to nearest USD 1.00.'),
    dict(num='08', category='COMMERCIAL CONTROLS', title='Tariff adjustment',
         summary='Derive a sell or cost rate from a named reference tariff, then apply controlled changes.',
         method='Reference uplift', behaviour='Result starts with the selected tariff outcome, then applies the approved adjustments in the stated order.',
         inputs=['Reference tariff', 'Base adjustment', 'Unit adjustment', 'Minimum or maximum', 'Order of operations'],
         safeguard='Show the tariff version, effective date and any missing-match warning. Never calculate silently from an unverified tariff.',
         example='Contract sea freight plus USD 15.00 per container and a 5 percent fuel uplift.'),
    dict(num='09', category='CUSTOMS', title='Customs clearance fee',
         summary='Price customs work from declaration type, fee structure, number of entries or number of invoice lines.',
         method='Declaration logic', behaviour='Result follows the selected declaration and line-count structure, subject to its cap.',
         inputs=['Base clearance fee', 'Entry type', 'Included entries or lines', 'Additional-entry rate', 'Maximum'],
         safeguard='Duties and taxes are separate from the service fee. Show the chosen declaration type and country-specific availability.',
         example='Import clearance: GBP 68.00 including one entry, GBP 14.00 per additional line, capped at GBP 180.00.'),
    dict(num='10', category='COMMERCIAL CONTROLS', title='Percentage uplift',
         summary='Apply a percentage to a defined base such as selected charges, goods value, insured value or customs value.',
         method='Percent of value', behaviour='Result = selected base x percentage, subject to the chosen floor or cap.',
         inputs=['Percentage', 'Calculation base', 'Minimum or maximum', 'Included charge scope'],
         safeguard='The calculation explanation must name every included amount. Operators should never have to infer the percentage base.',
         example='Cargo insurance: 0.38 percent of declared goods value, minimum GBP 25.00.'),
    dict(num='11', category='TRANSPORT', title='Distance zone pricing',
         summary='Price cartage using origin and destination zones, distance, equipment and optional volumetric conversion.',
         method='Zone + distance', behaviour='Result uses the matching zone pair and the selected distance or quantity bands.',
         inputs=['Origin and destination zones', 'Distance source', 'Vehicle or equipment', 'Base, bands and minimum'],
         safeguard='Identify the route and distance source in the result. A fallback zone must be explicit rather than hidden.',
         example='Bristol to London zone: GBP 95.00 base plus GBP 1.20 per road mile after 35 miles.'),
    dict(num='12', category='TRANSPORT', title='Transport band pricing',
         summary='Apply freight or transport rates by a measure such as chargeable weight, cbm, packages or equipment count.',
         method='Mode-aware bands', behaviour='Result follows the applicable transport band and its stated treatment.',
         inputs=['Transport mode', 'Measure source', 'Bands', 'Minimum, maximum and restrictions'],
         safeguard='Use warnings for out-of-range values, restricted bands and disagreement between rated and invoiced measures.',
         example='LCL freight by cbm with an origin-port restriction.'),
    dict(num='13', category='OPERATIONS', title='First plus additional',
         summary='Charge a different rate for the first item and each additional item.',
         method='First item + extras', behaviour='Result = first-item amount + max(item count - 1, 0) x additional-item amount.',
         inputs=['First-item amount', 'Additional-item amount', 'Item count', 'Item definition'],
         safeguard='Define the counted item clearly: consignment, document, line, pallet or package.',
         example='Delivery booking: GBP 18.00 first booking plus GBP 6.00 per additional booking.'),
    dict(num='14', category='PRESENTATION', title='Quote note',
         summary='Add a controlled explanatory note and optional amount to a customer quote without pretending it is an operational cost.',
         method='Text + amount', behaviour='Result adds the approved text line and optional amount to the selected quote output.',
         inputs=['Customer-facing description', 'Optional amount', 'Visibility setting'],
         safeguard='Clearly distinguish a presentation note from a calculated charge. Notes should not create supplier cost or margin automatically.',
         example='Optional weekend delivery cover - GBP 65.00.'),
    dict(num='15', category='WAREHOUSE', title='Package profile pricing',
         summary='Apply warehouse handling rates by package profile, such as full pallet, carton, case or loose unit.',
         method='Pack-aware', behaviour='Result converts the ordered quantity through the approved package hierarchy and applies the matching profile rates.',
         inputs=['Package profile', 'Conversion hierarchy', 'Rate per profile', 'Quantity source'],
         safeguard='Show the conversion path and any remainder so the result can be checked quickly on the warehouse floor.',
         example='Pick: GBP 2.10 per full pallet, GBP 0.34 per carton and GBP 0.08 per loose unit.'),
    dict(num='16', category='WAREHOUSE', title='Location profile pricing',
         summary='Apply a charge based on a warehouse location profile, regardless of the product held there.',
         method='Location-aware', behaviour='Result charges each qualifying location once for the configured billing period.',
         inputs=['Location profile', 'Rate', 'Billing period', 'Occupancy rule'],
         safeguard='Avoid double counting mixed-product locations. The guide should state whether a location was active during the period.',
         example='Temperature-controlled bay: GBP 14.00 per occupied location per week.'),
    dict(num='17', category='WAREHOUSE', title='Time period pricing',
         summary='Apply time-based storage or service charges by day, week, hour or another agreed period.',
         method='Storage time', behaviour='Result uses elapsed eligible time x the applicable rate or band.',
         inputs=['Start and end event', 'Time unit', 'Bands', 'Calendar exclusions'],
         safeguard='State which events start and stop the clock, and whether weekends and public holidays are excluded.',
         example='Storage: GBP 0.42 per pallet per day after two free days, excluding Sundays.'),
    dict(num='18', category='COMMERCIAL CONTROLS', title='Percentage bands',
         summary='Apply different percentage rules at defined value, weight or volume thresholds.',
         method='Tiered percent', behaviour='Result applies the percentage associated with the matching or progressive band.',
         inputs=['Break thresholds', 'Percentages', 'Calculation base', 'Target charge or group'],
         safeguard='Show whether bands are based on values or measures, and list the charges or groups affected.',
         example='Hazardous-goods supplement: 8 percent under GBP 2,000 value, 5 percent at GBP 2,000 and above.'),
    dict(num='19', category='FREIGHT PRICING', title='Best outcome rate',
         summary='Select the highest outcome from independently calculated minimum, weight, volume and fixed-price rules.',
         method='Highest charge wins', behaviour='Result = highest eligible outcome among the configured charge tests.',
         inputs=['Minimum rate', 'Weight rate', 'Volume rate', 'Flat amount', 'Multiples'],
         safeguard='Show every candidate calculation and clearly identify the winning one. Do not hide a lower alternative.',
         example='Air handling: maximum of GBP 60.00 minimum, GBP 0.35/kg, GBP 85.00/cbm or GBP 45.00 flat.'),
    dict(num='20', category='FINANCE', title='Disbursement finance charge',
         summary='Calculate a finance charge for duties, taxes or other disbursements funded before customer settlement.',
         method='Time + rate', behaviour='Result = eligible amount x effective annual rate x eligible days / 365.',
         inputs=['Prime rate', 'Uplift', 'Effective annual rate', 'Credit terms', 'Eligible days', 'Target amount'],
         safeguard='Use an approved finance-rate source and show the date, credit terms and whether only overdue days are included.',
         example='Duty advance: 9.5 percent annual rate over 21 eligible days.'),
    dict(num='21', category='COMMERCIAL CONTROLS', title='Value bands',
         summary='Apply rate bands based on a declared financial value such as goods, insurance, customs or invoice value.',
         method='Value-driven scale', behaviour='Result uses the selected value source and its matching band.',
         inputs=['Value source', 'Breaks', 'Rates', 'Minimum or maximum'],
         safeguard='The value source must be versioned and visible. Warn when the selected value has not been supplied or is stale.',
         example='Insurance administration: GBP 18.00 below GBP 10,000 goods value; GBP 32.00 at GBP 10,000 and above.'),
    dict(num='22', category='CONTRACT PRICING', title='Volume balance discount',
         summary='Apply an agreed discount when cumulative contracted volume reaches a defined performance threshold.',
         method='Period performance', behaviour='Result uses the achieved average or total contracted volume to determine whether the discount is eligible.',
         inputs=['Contract', 'Period', 'Pivot volume', 'Discount rate', 'Eligible jobs'],
         safeguard='Show the contract period, contributing shipments and whether the discount is provisional or final.',
         example='ULD air freight discount becomes eligible once monthly average chargeable weight reaches 1,500 kg.'),
    dict(num='23', category='OPERATIONS', title='Package count pricing',
         summary='Combine a base amount, rate per kg, first-package amount and additional-package amount.',
         method='Count + weight', behaviour='Result = base + weight element + first package + additional packages.',
         inputs=['Base amount', 'Rate per kg', 'First-package amount', 'Additional-package amount', 'Weight and count'],
         safeguard='Display each component in the result so users can check the selected count and weight sources.',
         example='Security screening: GBP 12.00 base, GBP 0.03/kg and GBP 1.50 per package after the first.'),
    dict(num='24', category='OPERATIONS', title='Release service pricing',
         summary='Apply an amount based on the release or collection service requested.',
         method='Release type', behaviour='Result uses the rate assigned to the selected release service.',
         inputs=['Release service', 'Rate', 'Eligibility conditions'],
         safeguard='Only show release services that are valid for the shipment, carrier and origin or destination process.',
         example='Express cargo release: GBP 35.00; standard release: GBP 15.00.'),
    dict(num='25', category='WAREHOUSE', title='Month split storage',
         summary='Use different storage rates before and after a defined point in the billing month.',
         method='Date threshold', behaviour='Result divides eligible storage time across the configured calendar threshold and applies each rate.',
         inputs=['Split day', 'Rate before split', 'Rate after split', 'Storage events'],
         safeguard='Show the exact date split and the days charged at each rate. This must be tied to the warehouse billing policy.',
         example='Pallet storage: GBP 0.38/day through day 14, then GBP 0.27/day for the remainder of the month.'),
]


def add_calculator_page(doc, calc):
    add_text(doc, f"CALCULATOR {calc['num']} / {calc['category']}", style='Guide Eyebrow')
    add_text(doc, calc['title'], style='Guide Heading 1')
    add_text(doc, calc['summary'], style='Guide Lead')
    add_table(doc, [('Pricing method', 'How it behaves'), (calc['method'], calc['behaviour'])],
              [5.8, 11.3], header=True, font_size=8.8)
    add_text(doc, '', size=2, after=5)
    add_two_panel(doc, calc['inputs'], calc['safeguard'])
    add_text(doc, '', size=2, after=5)
    add_example(doc, calc['example'])
    add_text(doc, 'Multideck interaction pattern', style='Guide Heading 2')
    add_text(doc,
             'Choose this pricing method from a plain-language selector. Multideck then reveals only the relevant inputs, '
             'validates missing data before rating, and keeps a live explanation beside the result. The charge remains '
             'independent on the supplier and customer sides unless an operator explicitly copies it.',
             style='Guide Body')


def add_standards_page(doc):
    add_text(doc, 'MODELLING STANDARDS', style='Guide Eyebrow')
    add_text(doc, 'What every Multideck calculator must carry', style='Guide Heading 1')
    add_table(doc, [
        ('Context', 'Organisation, legal entity, branch, trade lane, service, mode, equipment and the charge side where relevant.'),
        ('Commercial terms', 'Currency, tax treatment, rounding, minimum, maximum, effective dates, rate owner and approval status.'),
        ('Calculation basis', 'The precise source of weight, volume, distance, time, value, count, tariff or supplier cost.'),
        ('Outcome', 'Calculated amount, local/base amount, exchange rate, formula explanation and the winning band or constraint.'),
        ('Audit', 'Who created or changed the rule, when it was effective, which inputs were used and why any override occurred.'),
    ], [4.3, 12.8], header=False, body_fill=PALE_2, font_size=8.8)
    add_text(doc, 'Design direction', style='Guide Heading 2')
    add_text(doc,
             'This guide is a product foundation, not a request to expose every option at once. Multideck should '
             'introduce calculator methods progressively, prioritising the freight, transport and commercial controls '
             'that operators use most. Each new method should reuse the same result panel, audit contract and independent '
             'supplier/customer model.', style='Guide Body')


def remove_blank_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for part in (section.header, section.footer):
        for p in list(part.paragraphs):
            if not p.text and not p._p.xpath('.//w:drawing'):
                continue


def build():
    crop_assets()
    doc = Document()
    style_document(doc)
    add_cover(doc)

    intro_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(intro_section, top=1.30, bottom=0.62)
    add_header(intro_section)
    add_footer(intro_section, labelled=True)
    add_intro_page(doc)
    doc.add_page_break()
    add_catalogue_page(doc)

    calc_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(calc_section, top=0.60, bottom=0.60)
    calc_section.header.is_linked_to_previous = False
    calc_section.header.paragraphs[0].clear()
    add_footer(calc_section, labelled=False)
    for index, calc in enumerate(CALCULATORS):
        add_calculator_page(doc, calc)
        if index < len(CALCULATORS) - 1:
            doc.add_page_break()

    standards_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(standards_section, top=1.30, bottom=0.62)
    add_header(standards_section)
    add_footer(standards_section, labelled=True)
    add_standards_page(doc)

    settings = doc.settings._element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings.append(update_fields)
    update_fields.set(qn('w:val'), 'true')

    doc.core_properties.title = 'Multideck Calculator Guide Template'
    doc.core_properties.subject = 'Reusable Multideck pricing and product reference guide'
    doc.core_properties.author = 'Multideck'
    doc.core_properties.keywords = 'Multideck, calculator, freight, pricing, guide, template'
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build()
